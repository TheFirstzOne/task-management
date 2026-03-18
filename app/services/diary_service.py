"""
DiaryService — Business logic for daily work diary
"""

import os
from datetime import datetime
from typing import Dict, List, Optional
from sqlalchemy.orm import Session
from app.repositories.diary_repo import DiaryRepository
from app.models.diary import DiaryEntry


class DiaryService:

    def __init__(self, db: Session) -> None:
        self.repo = DiaryRepository(db)

    def create_entry(self, content: str) -> DiaryEntry:
        return self.repo.create(content)

    def get_all_entries(self) -> List[DiaryEntry]:
        return self.repo.get_all()

    def get_entries_grouped(self) -> Dict[str, List[DiaryEntry]]:
        return self.repo.get_grouped_by_date()

    def update_entry(self, entry_id: int, content: str) -> Optional[DiaryEntry]:
        return self.repo.update(entry_id, content)

    def delete_entry(self, entry_id: int) -> bool:
        return self.repo.delete(entry_id)

    def export_to_word(self, filepath: str) -> str:
        """Export all diary entries to a Word document. Returns filepath."""
        from docx import Document
        from docx.shared import Pt, Inches

        doc = Document()

        # Page setup A4
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        entries = self.repo.get_all()  # newest first

        if not entries:
            doc.add_paragraph("ยังไม่มีบันทึกการทำงาน")
            doc.save(filepath)
            return filepath

        for i, entry in enumerate(entries):
            if i > 0:
                doc.add_paragraph()
                doc.add_paragraph("=" * 80)

            # Date header
            date_str = entry.created_at.strftime("%d/%m/%Y %H:%M:%S")
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"\U0001f4c5 วันที่: {date_str}")
            date_run.bold = True
            date_run.font.size = Pt(12)

            # Content
            content_para = doc.add_paragraph()
            content_run = content_para.add_run(entry.content)
            content_run.font.size = Pt(11)

        doc.save(filepath)
        return filepath

    def export_to_pdf(self, filepath: str) -> str:
        """Export all diary entries to a PDF document. Returns filepath."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # ── Register Thai font (Tahoma) ──────────────────────────────
        _FONT = "Tahoma"
        for _fp in [r"C:\Windows\Fonts\tahoma.ttf",
                    "/usr/share/fonts/truetype/msttcorefonts/Tahoma.ttf"]:
            import os as _os
            if _os.path.exists(_fp):
                try:
                    pdfmetrics.registerFont(TTFont(_FONT, _fp))
                except Exception:
                    _FONT = "Helvetica"
                break
        else:
            _FONT = "Helvetica"

        # ── Styles (Blue-White theme) ────────────────────────────────
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DiaryTitle", fontName=_FONT, fontSize=18,
            textColor="#2563EB", spaceAfter=6,
        )
        date_style = ParagraphStyle(
            "DiaryDate", fontName=_FONT, fontSize=11,
            textColor="#2563EB", spaceBefore=10, spaceAfter=4,
            backColor="#EFF6FF", borderPad=4,
        )
        body_style = ParagraphStyle(
            "DiaryBody", fontName=_FONT, fontSize=11,
            textColor="#1E293B", leading=16, spaceAfter=6,
        )
        meta_style = ParagraphStyle(
            "DiaryMeta", fontName=_FONT, fontSize=9,
            textColor="#64748B", spaceAfter=2,
        )

        doc = SimpleDocTemplate(
            filepath, pagesize=A4,
            leftMargin=2*cm, rightMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm,
        )

        entries = self.repo.get_all()  # newest first
        story = []

        # Title
        story.append(Paragraph("บันทึกการทำงาน (Job Diary)", title_style))
        from datetime import datetime as _dt
        story.append(Paragraph(
            f"ออกรายงาน: {_dt.now().strftime('%d/%m/%Y %H:%M')}  |  ทั้งหมด {len(entries)} รายการ",
            meta_style,
        ))
        story.append(HRFlowable(width="100%", thickness=1, color="#2563EB", spaceAfter=12))

        if not entries:
            story.append(Paragraph("ยังไม่มีบันทึกการทำงาน", body_style))
        else:
            for entry in entries:
                date_str = entry.created_at.strftime("%d/%m/%Y  %H:%M น.")
                story.append(Paragraph(f"📅  {date_str}", date_style))
                # Replace newlines with <br/> for ReportLab
                safe_content = entry.content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
                story.append(Paragraph(safe_content, body_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color="#CBD5E1", spaceAfter=4))

        doc.build(story)
        return filepath
