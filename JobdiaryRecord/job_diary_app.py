import flet as ft
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re


class JobDiaryApp:
    def __init__(self, page: ft.Page):
        self.page = page
        self.page.title = "บันทึกการทำงานรายวัน"
        self.page.window.width = 900
        self.page.window.height = 700
        self.page.window.resizable = True
        self.page.padding = 20

        # ชื่อไฟล์ Word
        self.word_file = "job_diary.docx"

        # สร้าง UI
        self.setup_ui()

    def setup_ui(self):
        # หัวข้อ
        title = ft.Text(
            "บันทึกการทำงานรายวัน",
            size=28,
            weight=ft.FontWeight.BOLD,
            color=ft.Colors.BLUE_700,
        )

        # สร้าง Tab สำหรับบันทึกและอ่านย้อนหลัง
        tabs = ft.Tabs(
            selected_index=0,
            animation_duration=300,
            tabs=[
                ft.Tab(
                    text="✍️ บันทึกใหม่",
                    icon=ft.Icons.EDIT,
                    content=self.create_write_tab(),
                ),
                ft.Tab(
                    text="📖 อ่านย้อนหลัง",
                    icon=ft.Icons.HISTORY,
                    content=self.create_history_tab(),
                ),
            ],
            expand=1,
            on_change=self.on_tab_change,
        )

        # เพิ่มส่วนประกอบทั้งหมดลงในหน้า
        self.page.add(
            ft.Column(
                controls=[
                    title,
                    ft.Divider(height=20, color=ft.Colors.BLUE_200),
                    tabs,
                ],
                spacing=10,
                expand=True,
            )
        )

    def create_write_tab(self):
        """สร้าง Tab สำหรับบันทึก"""
        # กล่องข้อความสำหรับบันทึก
        self.diary_field = ft.TextField(
            label="บันทึกการทำงาน",
            hint_text="กรอกรายละเอียดการทำงานของคุณที่นี่...",
            multiline=True,
            min_lines=15,
            max_lines=20,
            border_color=ft.Colors.BLUE_400,
            focused_border_color=ft.Colors.BLUE_700,
            expand=True,
        )

        # ปุ่มบันทึก
        save_button = ft.ElevatedButton(
            "บันทึกลงไฟล์ Word",
            icon=ft.Icons.SAVE,
            on_click=self.save_to_word,
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.BLUE_700,
                color=ft.Colors.WHITE,
                padding=15,
            ),
        )

        # ปุ่มเคลียร์
        clear_button = ft.ElevatedButton(
            "ล้างข้อความ",
            icon=ft.Icons.CLEAR,
            on_click=self.clear_text,
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.ORANGE_700,
                color=ft.Colors.WHITE,
                padding=15,
            ),
        )

        # ปุ่มเปิดไฟล์
        open_button = ft.ElevatedButton(
            "เปิดไฟล์ Word",
            icon=ft.Icons.FOLDER_OPEN,
            on_click=self.open_word_file,
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.GREEN_700,
                color=ft.Colors.WHITE,
                padding=15,
            ),
        )

        # แถวปุ่ม
        button_row = ft.Row(
            controls=[save_button, clear_button, open_button],
            alignment=ft.MainAxisAlignment.CENTER,
            spacing=10,
        )

        # ข้อความสถานะ
        self.status_text = ft.Text(
            "",
            size=14,
            color=ft.Colors.GREEN_700,
            text_align=ft.TextAlign.CENTER,
        )

        # ข้อมูลไฟล์
        self.file_info = ft.Text(
            f"ไฟล์บันทึก: {self.word_file}",
            size=12,
            color=ft.Colors.GREY_700,
            italic=True,
        )

        return ft.Container(
            content=ft.Column(
                controls=[
                    self.diary_field,
                    ft.Container(height=10),
                    button_row,
                    ft.Container(height=10),
                    self.status_text,
                    self.file_info,
                ],
                spacing=10,
                expand=True,
            ),
            padding=10,
            expand=True,
        )

    def create_history_tab(self):
        """สร้าง Tab สำหรับอ่านบันทึกย้อนหลัง"""
        # พื้นที่แสดงรายการวันที่
        self.date_list = ft.ListView(
            spacing=10,
            padding=10,
            expand=True,
        )

        # พื้นที่แสดงเนื้อหาบันทึก
        self.history_content = ft.Column(
            spacing=10,
            scroll=ft.ScrollMode.AUTO,
            expand=True,
        )

        # ปุ่มรีเฟรช
        refresh_button = ft.IconButton(
            icon=ft.Icons.REFRESH,
            tooltip="โหลดข้อมูลใหม่",
            on_click=self.load_history,
            icon_color=ft.Colors.BLUE_700,
        )

        # หัวข้อรายการวันที่
        date_list_header = ft.Row(
            controls=[
                ft.Text(
                    "📅 รายการวันที่",
                    size=16,
                    weight=ft.FontWeight.BOLD,
                    color=ft.Colors.BLUE_700,
                ),
                refresh_button,
            ],
            alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
        )

        # คอลัมน์ซ้าย - รายการวันที่
        left_panel = ft.Container(
            content=ft.Column(
                controls=[
                    date_list_header,
                    ft.Divider(height=1, color=ft.Colors.GREY_400),
                    self.date_list,
                ],
                spacing=5,
            ),
            width=250,
            bgcolor=ft.Colors.GREY_100,
            border_radius=10,
            padding=10,
        )

        # คอลัมน์ขวา - เนื้อหาบันทึก
        right_panel = ft.Container(
            content=ft.Column(
                controls=[
                    ft.Text(
                        "📝 เนื้อหาบันทึก",
                        size=16,
                        weight=ft.FontWeight.BOLD,
                        color=ft.Colors.BLUE_700,
                    ),
                    ft.Divider(height=1, color=ft.Colors.GREY_400),
                    self.history_content,
                ],
                spacing=5,
            ),
            expand=True,
            bgcolor=ft.Colors.WHITE,
            border_radius=10,
            padding=10,
            border=ft.border.all(1, ft.Colors.GREY_300),
        )

        return ft.Container(
            content=ft.Row(
                controls=[left_panel, right_panel],
                spacing=10,
                expand=True,
            ),
            padding=10,
            expand=True,
        )

    def on_tab_change(self, e):
        """เมื่อเปลี่ยน Tab"""
        if e.control.selected_index == 1:  # Tab อ่านย้อนหลัง
            self.load_history(None)

    def load_history(self, e):
        """โหลดประวัติบันทึกจากไฟล์ Word"""
        self.date_list.controls.clear()
        self.history_content.controls.clear()

        if not os.path.exists(self.word_file):
            self.history_content.controls.append(
                ft.Container(
                    content=ft.Column(
                        controls=[
                            ft.Icon(ft.Icons.INFO_OUTLINE, size=64, color=ft.Colors.GREY_400),
                            ft.Text(
                                "ยังไม่มีบันทึก",
                                size=18,
                                color=ft.Colors.GREY_600,
                                weight=ft.FontWeight.BOLD,
                            ),
                            ft.Text(
                                "กรุณาบันทึกข้อมูลใน Tab 'บันทึกใหม่' ก่อน",
                                size=14,
                                color=ft.Colors.GREY_500,
                            ),
                        ],
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                        spacing=10,
                    ),
                    alignment=ft.alignment.center,
                    expand=True,
                )
            )
            self.page.update()
            return

        try:
            # อ่านไฟล์ Word
            doc = Document(self.word_file)

            # แยกบันทึกตามวันที่
            entries = self.parse_diary_entries(doc)

            if not entries:
                self.history_content.controls.append(
                    ft.Text(
                        "ไม่พบบันทึกในไฟล์",
                        size=16,
                        color=ft.Colors.GREY_600,
                    )
                )
            else:
                # แสดงรายการวันที่
                for date_key in entries.keys():
                    date_button = ft.Container(
                        content=ft.Row(
                            controls=[
                                ft.Icon(ft.Icons.CALENDAR_TODAY, size=16, color=ft.Colors.BLUE_600),
                                ft.Text(
                                    date_key,
                                    size=14,
                                    weight=ft.FontWeight.W_500,
                                ),
                            ],
                            spacing=8,
                        ),
                        padding=10,
                        border_radius=8,
                        bgcolor=ft.Colors.WHITE,
                        ink=True,
                        on_click=lambda e, d=date_key: self.show_date_entries(d, entries[d]),
                    )
                    self.date_list.controls.append(date_button)

                # แสดงวันที่แรกโดยอัตโนมัติ
                first_date = list(entries.keys())[0]
                self.show_date_entries(first_date, entries[first_date])

            self.page.update()

        except Exception as ex:
            self.history_content.controls.append(
                ft.Text(
                    f"เกิดข้อผิดพลาด: {str(ex)}",
                    size=14,
                    color=ft.Colors.RED_700,
                )
            )
            self.page.update()

    def parse_diary_entries(self, doc):
        """แยกบันทึกออกเป็นรายการตามวันที่"""
        entries = {}
        current_date = None
        current_time = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()

            # ตรวจสอบว่าเป็นบรรทัดวันที่หรือไม่
            date_match = re.search(r'📅 วันที่: (\d{2}/\d{2}/\d{4}) (\d{2}:\d{2}:\d{2})', text)

            if date_match:
                # บันทึกข้อมูลเก่าก่อน
                if current_date and current_content:
                    date_key = current_date
                    if date_key not in entries:
                        entries[date_key] = []
                    entries[date_key].append({
                        'time': current_time,
                        'content': '\n'.join(current_content)
                    })

                # เริ่มบันทึกใหม่
                current_date = date_match.group(1)
                current_time = date_match.group(2)
                current_content = []
            elif text and text != '=' * 80:
                # เพิ่มเนื้อหา
                if current_date:
                    current_content.append(text)

        # บันทึกข้อมูลสุดท้าย
        if current_date and current_content:
            date_key = current_date
            if date_key not in entries:
                entries[date_key] = []
            entries[date_key].append({
                'time': current_time,
                'content': '\n'.join(current_content)
            })

        # เรียงวันที่จากใหม่ไปเก่า
        sorted_entries = dict(sorted(entries.items(), reverse=True))
        return sorted_entries

    def show_date_entries(self, date, entries):
        """แสดงบันทึกของวันที่ที่เลือก"""
        self.history_content.controls.clear()

        # หัวข้อวันที่
        self.history_content.controls.append(
            ft.Container(
                content=ft.Text(
                    f"📅 {date}",
                    size=20,
                    weight=ft.FontWeight.BOLD,
                    color=ft.Colors.BLUE_700,
                ),
                padding=ft.padding.only(bottom=10),
            )
        )

        # แสดงบันทึกแต่ละรายการ
        for i, entry in enumerate(entries):
            card = ft.Card(
                content=ft.Container(
                    content=ft.Column(
                        controls=[
                            ft.Row(
                                controls=[
                                    ft.Icon(ft.Icons.ACCESS_TIME, size=16, color=ft.Colors.BLUE_300),
                                    ft.Text(
                                        entry['time'],
                                        size=14,
                                        weight=ft.FontWeight.BOLD,
                                        color=ft.Colors.BLUE_300,
                                    ),
                                ],
                                spacing=5,
                            ),
                            ft.Divider(height=1, color=ft.Colors.GREY_700),
                            ft.Text(
                                entry['content'],
                                size=14,
                                color=ft.Colors.WHITE,
                                selectable=True,
                            ),
                        ],
                        spacing=10,
                    ),
                    padding=15,
                    bgcolor=ft.Colors.BLUE_GREY_900,
                ),
                elevation=2,
            )
            self.history_content.controls.append(card)

        self.page.update()

    def save_to_word(self, e):
        # ตรวจสอบว่ามีข้อความหรือไม่
        if not self.diary_field.value or self.diary_field.value.strip() == "":
            self.show_status("กรุณากรอกข้อความก่อนบันทึก!", ft.Colors.RED_700)
            return

        try:
            # สร้างหรือเปิดไฟล์ Word
            if os.path.exists(self.word_file):
                doc = Document(self.word_file)
            else:
                doc = Document()
                # ตั้งค่าหน้ากระดาษ
                section = doc.sections[0]
                section.page_height = Inches(11.69)  # A4
                section.page_width = Inches(8.27)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)

            # เพิ่มวันที่และเวลา
            now = datetime.now()
            date_str = now.strftime("%d/%m/%Y %H:%M:%S")

            # เพิ่มเส้นแบ่งถ้าไม่ใช่บันทึกแรก
            if len(doc.paragraphs) > 0:
                doc.add_paragraph()
                doc.add_paragraph("=" * 80)

            # เพิ่มวันที่
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"📅 วันที่: {date_str}")
            date_run.bold = True
            date_run.font.size = Pt(12)
            date_run.font.color.rgb = None  # ใช้สีดำ

            # เพิ่มเนื้อหาบันทึก
            content_para = doc.add_paragraph()
            content_run = content_para.add_run(self.diary_field.value)
            content_run.font.size = Pt(11)

            # บันทึกไฟล์
            doc.save(self.word_file)

            # แสดงข้อความสำเร็จ
            self.show_status(f"✓ บันทึกสำเร็จ! เพิ่มลงไฟล์ {self.word_file}", ft.Colors.GREEN_700)

            # ล้างข้อความในกล่อง
            self.diary_field.value = ""
            self.page.update()

        except Exception as ex:
            self.show_status(f"เกิดข้อผิดพลาด: {str(ex)}", ft.Colors.RED_700)

    def clear_text(self, e):
        self.diary_field.value = ""
        self.status_text.value = ""
        self.page.update()

    def open_word_file(self, e):
        if os.path.exists(self.word_file):
            try:
                os.startfile(self.word_file)  # สำหรับ Windows
                self.show_status(f"กำลังเปิดไฟล์ {self.word_file}...", ft.Colors.BLUE_700)
            except Exception as ex:
                self.show_status(f"ไม่สามารถเปิดไฟล์: {str(ex)}", ft.Colors.RED_700)
        else:
            self.show_status("ยังไม่มีไฟล์บันทึก กรุณาบันทึกข้อมูลก่อน", ft.Colors.ORANGE_700)

    def show_status(self, message, color):
        self.status_text.value = message
        self.status_text.color = color
        self.page.update()


def main(page: ft.Page):
    JobDiaryApp(page)


if __name__ == "__main__":
    ft.app(target=main)
